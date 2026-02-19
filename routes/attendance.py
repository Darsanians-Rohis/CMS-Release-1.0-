from datetime import datetime, timezone, timedelta
from io import BytesIO
from flask import Blueprint, request, jsonify, Response
from flask_login import login_required, current_user
from sqlalchemy.exc import IntegrityError
from docx import Document
from extensions import db
from models import Session, Attendance, User
from serializers import serialize_attendance, serialize_user
from utils import can_mark_attendance, is_core_user

bp = Blueprint("attendance", __name__)

WIB = timezone(timedelta(hours=7))
ADMIN_ROLES = {"admin", "ketua", "pembina"}


def _require_admin():
    if current_user.role not in ADMIN_ROLES:
        return jsonify({"success": False, "message": "Access denied"}), 403


def _record_attendance(session_id, user_id, status, attendance_type):
    """Core insert logic shared between regular and core attendance."""
    s = Session.query.get(session_id)
    if not s:
        return jsonify({"success": False, "error": "not_found", "message": "Session not found"}), 404
    if s.is_locked:
        return jsonify({"success": False, "error": "session_locked", "message": "Session is locked"}), 403

    existing = Attendance.query.filter_by(
        session_id=session_id, user_id=user_id, attendance_type=attendance_type
    ).first()
    if existing:
        return jsonify({"success": False, "error": "already_marked", "message": "Attendance already recorded"}), 409

    att = Attendance(
        session_id=session_id,
        user_id=user_id,
        status=status,
        attendance_type=attendance_type,
        timestamp=datetime.now(WIB),
    )
    try:
        db.session.add(att)
        db.session.commit()
        return jsonify({"success": True, "attendance": serialize_attendance(att)}), 201
    except IntegrityError:
        db.session.rollback()
        return jsonify({"success": False, "error": "already_marked", "message": "Attendance already recorded"}), 409
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": "database_error", "message": str(e)}), 500


@bp.route("/api/attendance", methods=["POST"])
@login_required
def api_attendance():
    data = request.get_json() or {}
    session_id = data.get("session_id")
    user_id = data.get("user_id")
    status = data.get("status")

    if not all([session_id, user_id, status]):
        return jsonify({"success": False, "error": "invalid_data", "message": "Missing required fields"}), 400

    try:
        session_id, user_id = int(session_id), int(user_id)
    except (ValueError, TypeError):
        return jsonify({"success": False, "error": "invalid_data", "message": "Invalid ID format"}), 400

    s = Session.query.get(session_id)
    if s and not can_mark_attendance(current_user, s.pic_id):
        return jsonify({"success": False, "error": "forbidden", "message": "No permission to mark attendance"}), 403

    return _record_attendance(session_id, user_id, status, "regular")


@bp.route("/api/attendance/core", methods=["POST"])
@login_required
def api_attendance_core():
    if not is_core_user(current_user):
        return jsonify({"success": False, "error": "forbidden", "message": "Access denied"}), 403

    data = request.get_json() or {}
    session_id = data.get("session_id")
    user_id = data.get("user_id")
    status = data.get("status")

    if not all([session_id, user_id, status]):
        return jsonify({"success": False, "error": "invalid_data", "message": "Missing required fields"}), 400

    try:
        session_id, user_id = int(session_id), int(user_id)
    except (ValueError, TypeError):
        return jsonify({"success": False, "error": "invalid_data", "message": "Invalid ID format"}), 400

    target = User.query.get(user_id)
    if not target or not is_core_user(target):
        return jsonify({"success": False, "error": "not_core_user", "message": "User is not a core member"}), 400

    return _record_attendance(session_id, user_id, status, "core")


@bp.route("/api/attendance/history")
@login_required
def attendance_history():
    records = Attendance.query.filter_by(user_id=current_user.id).all()
    summary = {
        "present": sum(1 for r in records if r.status == "present"),
        "absent": sum(1 for r in records if r.status == "absent"),
        "excused": sum(1 for r in records if r.status == "excused"),
        "late": sum(1 for r in records if r.status == "late"),
        "total": len(records),
    }
    return jsonify({"success": True, "records": [serialize_attendance(r) for r in records], "summary": summary})


@bp.route("/api/attendance/history/all")
@login_required
def attendance_history_all():
    err = _require_admin()
    if err:
        return err
    users = User.query.filter_by(role="member").order_by(User.name).all()
    return jsonify({"success": True, "members": [serialize_user(u) for u in users]})


@bp.route("/api/attendance/history/<int:user_id>")
@login_required
def attendance_history_for_user(user_id):
    if current_user.role not in ADMIN_ROLES and current_user.id != user_id:
        return jsonify({"success": False, "message": "Access denied"}), 403

    user = User.query.get_or_404(user_id)
    records = Attendance.query.filter_by(user_id=user_id).all()
    summary = {
        "present": sum(1 for r in records if r.status == "present"),
        "absent": sum(1 for r in records if r.status == "absent"),
        "excused": sum(1 for r in records if r.status == "excused"),
        "late": sum(1 for r in records if r.status == "late"),
        "total": len(records),
    }
    return jsonify({
        "success": True,
        "user": serialize_user(user),
        "records": [serialize_attendance(r) for r in records],
        "summary": summary,
    })


@bp.route("/api/export/attendance/<int:session_id>")
@login_required
def export_attendance(session_id):
    err = _require_admin()
    if err:
        return err

    s = Session.query.get_or_404(session_id)
    records = (
        db.session.query(Attendance, User.name, User.email, User.role)
        .join(User, Attendance.user_id == User.id)
        .filter(Attendance.session_id == session_id)
        .order_by(User.name)
        .all()
    )

    if not records:
        return jsonify({"success": False, "message": "No attendance records found"}), 404

    doc = Document()
    doc.add_heading(f"Attendance Report: {s.name}", 0)
    doc.add_paragraph(f"Date: {s.date}")
    doc.add_paragraph(f"Total Attendees: {len(records)}")
    doc.add_paragraph("")

    summary = {k: sum(1 for a, *_ in records if a.status == k) for k in ("present", "absent", "excused", "late")}
    doc.add_heading("Summary", level=1)
    st = doc.add_table(rows=5, cols=2)
    st.style = "Light Grid Accent 1"
    for i, (label, val) in enumerate([
        ("Status", "Count"),
        ("Present", str(summary["present"])),
        ("Absent", str(summary["absent"])),
        ("Excused", str(summary["excused"])),
        ("Late", str(summary["late"])),
    ]):
        st.rows[i].cells[0].text = label
        st.rows[i].cells[1].text = val

    doc.add_paragraph("")
    doc.add_heading("Detailed Records", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Name", "Role", "Status", "Time", "Type"]):
        hdr[i].text = h

    for att, name, email, role in records:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = role.capitalize()
        row[2].text = att.status.capitalize()
        row[3].text = att.timestamp.astimezone(WIB).strftime("%H:%M") if att.timestamp else ""
        row[4].text = att.attendance_type.capitalize()

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"attendance_{s.name.replace(' ', '_')}_{s.date}.docx"
    return Response(
        bio,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
